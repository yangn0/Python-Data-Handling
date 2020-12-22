from docx import Document
import os.path

'''
document.add_heading('Document Title', 1)
p = document.add_paragraph('A plain paragraph having some ')
document.save('demo.docx')'''


def file_extension(path): 
    return os.path.splitext(path)[1] 

def input_code(path):
    global q
    c=os.getcwd()
    os.chdir(path)
    cwd=os.getcwd()
    l=os.listdir()
    for i in l:
        if(os.path.isfile(i)):
            if(file_extension(i) in a):
                os.chdir(cwd)
                document.add_heading(i, 1)
                q+=1
                try:
                    print(i)
                    with open(i,encoding='UTF-8') as f:
                        for u in f:
                            if (u.rstrip()!=''):
                                document.add_paragraph(u.rstrip())
                    document.add_paragraph('')
                except:
                    print(i+'编码有误，请手动添加')
        if(os.path.isdir(i)):
            input_code(i)
    os.chdir(c)
    

''' k=list()
    for i in l:
        if(os.path.isdir(i)):
            k.append(1)
        if(os.path.isfile(i)):
            k.append(0)
    if 1 not in k:
        k=list()
        return 1'''

q=0
path=input('文件夹：')
a=input('后缀：').split(' ')
document = Document()
n_path=os.getcwd()
input_code(path)
os.chdir(n_path)
document.save('1.docx')
print(q)
input("汇总成功，回车退出,已保存为1.docx。")
